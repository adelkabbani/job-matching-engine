"""
Document Generation Service.
Uses python-docx to create .docx files for CVs and Cover Letters.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from typing import Dict, List

def generate_cv_docx(cv_data: Dict, output_path: str):
    """Generates a professional CV docx file."""
    doc = Document()
    
    # 1. Header (Name & Contact)
    name = cv_data.get('full_name', 'Your Name')
    header = doc.add_heading(name, 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_info = []
    if cv_data.get('email'): contact_info.append(cv_data['email'])
    if cv_data.get('phone'): contact_info.append(cv_data['phone'])
    if cv_data.get('location'): contact_info.append(cv_data['location'])
    
    contact_p = doc.add_paragraph(' | '.join(contact_info))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Summary
    if cv_data.get('summary'):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(cv_data['summary'])
    
    # 3. Skills
    if cv_data.get('skills'):
        doc.add_heading('Technical Skills', level=1)
        skills_text = ', '.join(cv_data['skills'])
        doc.add_paragraph(skills_text)
    
    # 4. Experience
    if cv_data.get('work_experience'):
        doc.add_heading('Work Experience', level=1)
        for exp in cv_data['work_experience']:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp.get('job_title', 'Title')} | {exp.get('company', 'Company')}")
            run.bold = True
            p.add_run(f"\n{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}")
            
            bullets = exp.get('description', [])
            if isinstance(bullets, str):
                bullets = [b.strip() for b in bullets.split('\n') if b.strip()]
            
            for bullet in bullets:
                doc.add_paragraph(bullet, style='List Bullet')
    
    # 5. Education
    if cv_data.get('education'):
        doc.add_heading('Education', level=1)
        for edu in cv_data['education']:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', 'Degree')} in {edu.get('field', 'Field')}")
            run.bold = True
            p.add_run(f"\n{edu.get('institution', 'University')} | {edu.get('end_date', 'Date')}")
    
    # Save
    doc.save(output_path)
    return output_path

def generate_cover_letter_docx(content: str, user_name: str, output_path: str):
    """Generates a professional cover letter docx."""
    doc = Document()
    
    # Header
    doc.add_heading(user_name, 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    
    doc.add_paragraph("\n")
    
    # Content
    for paragraph in content.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
            
    # Sign off
    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph(user_name)
    
    doc.save(output_path)
    return output_path
