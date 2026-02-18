from docx import Document
from pathlib import Path
import sys

# Setup logging
log_path = Path(__file__).parent / ".tmp" / "create_cv.log"
def log(msg):
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(msg + "\n")

def create_dummy_cv():
    log("Starting create_dummy_cv...")
    try:
        doc = Document()
        
        doc.add_heading('John Doe', 0)
        
        doc.add_paragraph('Email: john.doe@example.com | Phone: +1-555-010-9988')
        doc.add_paragraph('Berlin, Germany')
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph('Experienced Software Engineer with a focus on Python and React.')
        
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('Languages: Python, JavaScript, C++')
        doc.add_paragraph('Frameworks: React, Django, Node.js')
        doc.add_paragraph('Tools: Docker, Git, AWS')
        
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Master of Computer Science')
        doc.add_paragraph('Technical University of Munich')
        
        doc.add_heading('Work Experience', level=1)
        doc.add_paragraph('Senior Software Engineer | TechCorp | 2020 - Present')
        doc.add_paragraph('Developed scalable microservices using Python and FastAPI.')
        
        output_path = Path(__file__).parent / ".tmp" / "dummy_cv.docx"
        doc.save(str(output_path))
        log(f"Created {output_path}")
    except Exception as e:
        log(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    create_dummy_cv()

