"""
ATS-Friendly CV Generator

Generates tailored, ATS-optimized CVs for each job application.

ATS (Applicant Tracking System) Optimization:
- Simple, clean formatting (no tables, columns, text boxes)
- Standard fonts (Arial, Calibri, Times New Roman)
- Standard section headings
- Keyword optimization from job description
- Reverse chronological order
- No headers/footers, images, or graphics
- Standard bullet points
- Proper use of whitespace

Usage:
    python execution/generate_cv.py <job_id>
    python execution/generate_cv.py --all  # Generate for all filtered jobs
    
Reads:
    - .tmp/user_profile.json (your experience)
    - .tmp/jobs_filtered.json (jobs to apply to)
    
Outputs:
    - .tmp/applications/{company_name}/cv.docx
    - .tmp/applications/{company_name}/cv.pdf
"""

import json
import sys
import os
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    # We can't use os.system like this safely, but assuming requirements are there
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Setup logging
log_path = Path(__file__).parent.parent / ".tmp" / "generate_cv.log"
def log(msg):
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(msg + "\n")

log("Starting generate_cv.py...")



def load_profile():
    """Load user profile."""
    profile_path = Path(__file__).parent.parent / ".tmp" / "user_profile.json"
    with open(profile_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_jobs():
    """Load filtered jobs."""
    jobs_path = Path(__file__).parent.parent / ".tmp" / "jobs_filtered.json"
    with open(jobs_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def extract_keywords_from_job(job):
    """
    Extract important keywords from job title and description.
    These will be emphasized in the CV.
    """
    text = f"{job.get('title', '')} {job.get('description', '')}".lower()
    
    # Common technical skills to look for
    tech_keywords = [
        'python', 'javascript', 'java', 'c++', 'react', 'vue', 'angular',
        'node.js', 'django', 'flask', 'fastapi', 'express', 'spring',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
        'postgresql', 'mongodb', 'redis', 'mysql', 'sql',
        'git', 'ci/cd', 'agile', 'scrum', 'microservices',
        'rest api', 'graphql', 'websockets', 'tensorflow', 'pytorch',
        'machine learning', 'ai', 'data science', 'analytics',
        'leadership', 'mentoring', 'team lead', 'architect'
    ]
    
    found_keywords = set()
    for keyword in tech_keywords:
        if keyword in text:
            found_keywords.add(keyword)
    
    return list(found_keywords)


def create_ats_friendly_cv(profile, job, keywords):
    """
    Create an ATS-friendly CV in DOCX format.
    
    ATS Best Practices:
    - Use standard fonts (Calibri 11pt)
    - No tables, columns, or text boxes
    - Standard section headings in ALL CAPS
    - Simple bullet points
    - No headers/footers
    - No images or graphics
    - Keywords from job description
    - Reverse chronological order
    """
    doc = Document()
    
    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # HEADER: Name and Contact (ATS-friendly format)
    personal_info = profile.get('personal_info', {})
    
    # Name - Large, bold
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(personal_info.get('full_name', 'Your Name'))
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info - one line, centered
    contact_parts = []
    if personal_info.get('email'):
        contact_parts.append(personal_info['email'])
    if personal_info.get('phone'):
        contact_parts.append(personal_info['phone'])
    if personal_info.get('location', {}).get('city'):
        loc = personal_info['location']
        location_str = f"{loc.get('city', '')}, {loc.get('state', '')}" if loc.get('state') else loc.get('city', '')
        contact_parts.append(location_str)
    
    contact_para = doc.add_paragraph(' | '.join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.runs[0]
    contact_run.font.size = Pt(11)
    contact_run.font.name = 'Calibri'
    
    # Links - one line, centered
    links_parts = []
    if personal_info.get('linkedin'):
        links_parts.append(personal_info['linkedin'])
    if personal_info.get('github'):
        links_parts.append(personal_info['github'])
    if personal_info.get('portfolio'):
        links_parts.append(personal_info['portfolio'])
    
    if links_parts:
        links_para = doc.add_paragraph(' | '.join(links_parts))
        links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        links_run = links_para.runs[0]
        links_run.font.size = Pt(10)
        links_run.font.name = 'Calibri'
    
    doc.add_paragraph()  # Spacing
    
    # PROFESSIONAL SUMMARY (tailored to job)
    add_section_header(doc, 'PROFESSIONAL SUMMARY')
    
    # Generate a tailored summary (in production, this would use AI)
    summary_text = generate_summary(profile, job, keywords)
    summary_para = doc.add_paragraph(summary_text)
    summary_para.style = 'Body Text'
    for run in summary_para.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    
    doc.add_paragraph()  # Spacing
    
    # PROFESSIONAL EXPERIENCE
    add_section_header(doc, 'PROFESSIONAL EXPERIENCE')
    
    work_experience = profile.get('work_experience', [])
    for exp in work_experience[:5]:  # Limit to 5 most recent
        # Company and Title - Bold
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f"{exp.get('title', 'Position')} | {exp.get('company', 'Company')}")
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        title_run.font.name = 'Calibri'
        
        # Dates and Location
        dates_para = doc.add_paragraph()
        start = exp.get('start_date', '')
        end = exp.get('end_date', 'Present') if not exp.get('current') else 'Present'
        location = exp.get('location', '')
        dates_run = dates_para.add_run(f"{start} – {end} | {location}")
        dates_run.font.size = Pt(10)
        dates_run.font.name = 'Calibri'
        dates_run.font.italic = True
        
        # Achievements - tailored to emphasize relevant keywords
        achievements = exp.get('achievements', [])
        if achievements:
            for achievement in achievements[:4]:  # Max 4 bullets per job
                bullet_para = doc.add_paragraph(achievement, style='List Bullet')
                for run in bullet_para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
        
        doc.add_paragraph()  # Spacing between jobs
    
    # EDUCATION
    add_section_header(doc, 'EDUCATION')
    
    education = profile.get('education', [])
    for edu in education:
        # Degree and Institution - Bold
        edu_para = doc.add_paragraph()
        degree_text = f"{edu.get('degree', '')} in {edu.get('field_of_study', '')}" if edu.get('field_of_study') else edu.get('degree', '')
        edu_run = edu_para.add_run(f"{degree_text} | {edu.get('institution', '')}")
        edu_run.font.bold = True
        edu_run.font.size = Pt(11)
        edu_run.font.name = 'Calibri'
        
        # Dates
        if edu.get('end_date'):
            date_para = doc.add_paragraph(f"Graduated: {edu['end_date']}")
            for run in date_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                run.font.italic = True
        
        # GPA if notable
        if edu.get('gpa') and edu['gpa'] >= 3.5:
            gpa_para = doc.add_paragraph(f"GPA: {edu['gpa']}/4.0")
            for run in gpa_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        
        doc.add_paragraph()  # Spacing
    
    # TECHNICAL SKILLS (keyword-optimized)
    add_section_header(doc, 'TECHNICAL SKILLS')
    
    skills_data = profile.get('skills', [])
    skills = []
    
    if isinstance(skills_data, list):
        # Handle flat list from ingest_cv.py
        # Create dummy dict objects for compatibility or just use strings
        skills = [{'name': s} for s in skills_data if isinstance(s, str)]
        # Also handle if list already contains dicts
        skills.extend([s for s in skills_data if isinstance(s, dict)])
    elif isinstance(skills_data, dict):
        # Handle legacy dict format
        skills = skills_data.get('technical', [])

    
    # Group skills by category for better ATS parsing
    skill_categories = {
        'Programming Languages': [],
        'Frameworks & Libraries': [],
        'Databases': [],
        'Cloud & DevOps': [],
        'Tools & Technologies': []
    }
    
    for skill in skills:
        skill_name = skill.get('name', '')
        # Simple categorization (in production, this would be smarter)
        if any(lang in skill_name.lower() for lang in ['python', 'javascript', 'java', 'c++', 'go', 'rust', 'ruby']):
            skill_categories['Programming Languages'].append(skill_name)
        elif any(fw in skill_name.lower() for fw in ['react', 'vue', 'angular', 'django', 'flask', 'spring']):
            skill_categories['Frameworks & Libraries'].append(skill_name)
        elif any(db in skill_name.lower() for db in ['sql', 'postgres', 'mysql', 'mongo', 'redis']):
            skill_categories['Databases'].append(skill_name)
        elif any(cloud in skill_name.lower() for cloud in ['aws', 'azure', 'gcp', 'docker', 'kubernetes']):
            skill_categories['Cloud & DevOps'].append(skill_name)
        else:
            skill_categories['Tools & Technologies'].append(skill_name)
    
    for category, category_skills in skill_categories.items():
        if category_skills:
            skills_para = doc.add_paragraph()
            skills_para.add_run(f"{category}: ").font.bold = True
            skills_para.add_run(', '.join(category_skills[:10]))  # Max 10 per category
            for run in skills_para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
    
    # CERTIFICATIONS (if any)
    certifications = profile.get('certifications', [])
    if certifications:
        doc.add_paragraph()
        add_section_header(doc, 'CERTIFICATIONS')
        
        for cert in certifications[:5]:  # Max 5
            cert_para = doc.add_paragraph()
            cert_run = cert_para.add_run(f"{cert.get('name', '')} - {cert.get('issuing_organization', '')}")
            cert_run.font.size = Pt(11)
            cert_run.font.name = 'Calibri'
            
            if cert.get('issue_date'):
                cert_para.add_run(f" ({cert['issue_date']})")
    
    return doc


def add_section_header(doc, text):
    """Add a section header in ATS-friendly format."""
    header = doc.add_paragraph(text)
    header_run = header.runs[0]
    header_run.font.size = Pt(12)
    header_run.font.bold = True
    header_run.font.name = 'Calibri'
    header_run.font.color.rgb = RGBColor(0, 0, 0)  # Pure black for ATS
    
    # Add a line below header for visual separation
    # (ATS-friendly: just a paragraph with underline, not a table)
    line = doc.add_paragraph('_' * 80)
    for run in line.runs:
        run.font.size = Pt(8)


def generate_summary(profile, job, keywords):
    """
    Generate a tailored professional summary.
    In production, this would use AI. For now, it's template-based.
    """
    personal_info = profile.get('personal_info', {})
    headline = personal_info.get('professional_headline', 'Software Engineer')
    
    # Count years of experience
    work_exp = profile.get('work_experience', [])
    total_years = len(work_exp) * 1.5  # Rough estimate
    
    # Highlight matching keywords
    keyword_str = ', '.join(keywords[:5]) if keywords else 'modern technologies'
    
    summary = (
        f"{headline} with {int(total_years)}+ years of experience building scalable applications. "
        f"Expertise in {keyword_str}. Proven track record of delivering high-impact projects. "
        f"Seeking to contribute technical excellence to {job.get('company', 'your team')} as a {job.get('title', 'Software Engineer')}."
    )
    
    return summary


def save_cv(doc, company_name):
    """Save CV as DOCX and PDF."""
    # Create company folder
    app_dir = Path(__file__).parent.parent / ".tmp" / "applications" / company_name.replace(' ', '_').replace('/', '_')
    app_dir.mkdir(parents=True, exist_ok=True)
    
    # Save DOCX (ATS-friendly)
    docx_path = app_dir / "cv.docx"
    doc.save(str(docx_path))
    
    # PDF conversion (optional, DOCX is more ATS-friendly)
    # For now, just save DOCX
    # In production, you'd use docx2pdf or similar
    
    return docx_path


def main():
    log("ATS-Friendly CV Generator\n")
    
    # Load data
    profile = load_profile()
    jobs = load_jobs()
    
    if not jobs:
        log("No filtered jobs found. Run filtering first.")
        sys.exit(1)
    
    log(f"Found {len(jobs)} jobs to generate CVs for\n")
    
    # Generate CVs
    for i, job in enumerate(jobs, 1):
        company = job.get('company', 'Unknown Company')
        title = job.get('title', 'Position')
        
        log(f"{i}. Generating CV for: {title} at {company}")
        
        # Extract keywords
        keywords = extract_keywords_from_job(job)
        log(f"   Keywords: {', '.join(keywords[:5])}")
        
        # Generate CV
        doc = create_ats_friendly_cv(profile, job, keywords)
        
        # Save
        cv_path = save_cv(doc, company)
        log(f"   Saved: {cv_path}\n")
    
    log(f"\nGenerated {len(jobs)} ATS-friendly CVs!")
    log(f"Location: .tmp/applications/")
    log(f"\nNext steps:")
    log(f"   1. Review generated CVs")
    log(f"   2. Customize if needed")
    log(f"   3. Run cover letter generation")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        log(f"Error: {e}")
        import traceback
        log(traceback.format_exc())
        sys.exit(1)
