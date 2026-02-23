import sys
import os
import json
import re
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Setup file logging
log_path = Path(__file__).parent.parent / ".tmp" / "ingest_cv.log"
def log(msg):
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(msg + "\n")
    print(msg)

log("Starting ingest_cv.py...")

try:
    import pdfplumber
    import docx
    # Import LLM service
    sys.path.append(str(Path(__file__).parent.parent))
    from backend.services.llm import extract_resume_data
    log("Imports successful.")
except Exception as e:
    log(f"Import Error: {e}")
    sys.exit(1)


def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        log(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        log(f"Error reading DOCX: {e}")
    return text

def ingest_cv(file_path):
    path = Path(file_path)
    if not path.exists():
        log(f"Error: File not found: {file_path}")
        return

    log(f"Parsing {file_path}...")
    
    if path.suffix.lower() == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif path.suffix.lower() == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        log("Unsupported file format. Please use PDF or DOCX.")
        return

    if not text:
        log("Failed to extract text.")
        return

    # Call LLM for parsing
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        log("Error: OPENROUTER_API_KEY not found in .env")
        return

    log("Extracting structured data using LLM...")
    try:
        resume_data = extract_resume_data(text, api_key)
        log("✅ LLM Extraction Successful.")
    except Exception as e:
        log(f"❌ LLM Parsing Failed: {e}")
        return

    # Load existing profile
    profile_path = Path(__file__).parent.parent / ".tmp" / "user_profile.json"
    if profile_path.exists():
        with open(profile_path, 'r', encoding='utf-8') as f:
            try:
                profile = json.load(f)
            except:
                profile = {}
    else:
        profile = {}

    # Merge Logic
    # 1. Update Personal Info
    if "personal_info" not in profile: profile["personal_info"] = {}
    
    new_info = resume_data.get("personal_info", {})
    # Only update empty fields to preserve manual edits if any
    for k, v in new_info.items():
        if v and not profile["personal_info"].get(k):
             profile["personal_info"][k] = v

    # 2. Update Skills (Union)
    current_skills = set(profile.get("skills", []))
    new_skills = set(resume_data.get("skills", {}).get("technical", []))
    # Add other skill categories if present
    new_skills.update(resume_data.get("skills", {}).get("soft", []))
    new_skills.update(resume_data.get("skills", {}).get("tools", []))
    
    current_skills.update(new_skills)
    profile["skills"] = list(current_skills)

    # 3. Update Education (Append unique)
    if "education" not in profile: profile["education"] = []
    new_education = resume_data.get("education", [])
    
    # Simple check to avoid exact duplicates
    existing_edu_str = [json.dumps(e, sort_keys=True) for e in profile["education"]]
    for edu in new_education:
        if json.dumps(edu, sort_keys=True) not in existing_edu_str:
             profile["education"].append(edu)

    # 4. Update Experience
    if "work_experience" not in profile: profile["work_experience"] = []
    # Similar append logic could go here, for now just replace if empty or append
    if resume_data.get("work_experience"):
         # For simplicity in this phase, we append new found ones
         profile["work_experience"].extend(resume_data["work_experience"])

    # 5. Projects
    if "projects" not in profile: profile["projects"] = []
    if resume_data.get("projects"):
        profile["projects"].extend(resume_data["projects"])

    # 6. Metadata
    profile["experience_level"] = resume_data.get("experience_level", "Entry Level")

    # Save
    with open(profile_path, 'w', encoding='utf-8') as f:
        json.dump(profile, f, indent=2, ensure_ascii=False)
        
    log(f"\n✅ Profile updated at {profile_path}")
    log(f"Skills Count: {len(profile['skills'])}")
    log(f"Experience Level: {profile['experience_level']}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        log("Usage: python execution/ingest_cv.py <path_to_cv>")
        sys.exit(1)
    
    ingest_cv(sys.argv[1])

def log(msg):
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(msg + "\n")

log("Starting ingest_cv.py...")

try:
    import pdfplumber
    import docx
    log("Imports successful.")
except Exception as e:
    log(f"Import Error: {e}")
    sys.exit(1)


def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text

def extract_contact_info(text):
    """Extract email and phone number."""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'(\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}'
    
    email = re.search(email_pattern, text)
    phone = re.search(phone_pattern, text)
    
    return {
        "email": email.group(0) if email else None,
        "phone": phone.group(0) if phone else None
    }

def extract_skills(text):
    """Extract skills based on a predefined keyword list."""
    # Common tech skills taxonomy
    skill_keywords = [
        "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Rust", "PHP", "Ruby",
        "React", "Angular", "Vue", "Node.js", "Django", "Flask", "FastAPI", "Spring", ".NET",
        "HTML", "CSS", "SQL", "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch",
        "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "Jenkins", "Git", "Linux",
        "Machine Learning", "Deep Learning", "NLP", "Pandas", "NumPy", "PyTorch", "TensorFlow",
        "Agile", "Scrum", "Jira", "CI/CD", "TDD"
    ]
    
    found_skills = []
    text_lower = text.lower()
    
    for skill in skill_keywords:
        # Simple word boundary check to avoid partial matches (e.g., "Go" in "Good")
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, text_lower):
            found_skills.append(skill)
            
    return found_skills

def extract_education(text):
    """Heuristic extraction of education."""
    education = []
    text_lower = text.lower()
    lines = text.split('\n')
    
    degree_keywords = ["bachelor", "master", "phd", "b.sc", "m.sc", "mba", "b.a", "m.a", "associate degree"]
    uni_keywords = ["university", "college", "institute", "school"]
    
    for i, line in enumerate(lines):
        line_lower = line.lower()
        
        # Check if line contains a degree
        found_degree = next((kw for kw in degree_keywords if kw in line_lower), None)
        
        if found_degree:
            # Found a potential degree line
            # Look for university in this line or nearby lines
            uni_name = "Unknown University"
            
            # Check current line for university
            if any(ukw in line_lower for ukw in uni_keywords):
                uni_name = line.strip()
            else:
                # Check next line
                if i + 1 < len(lines):
                    next_line = lines[i+1]
                    if any(ukw in next_line.lower() for ukw in uni_keywords):
                        uni_name = next_line.strip()
            
            # Clean up degree name (simple heuristic)
            degree_name = line.strip()
            
            # Avoid duplicates
            if not any(e['degree'] == degree_name for e in education):
                education.append({
                    "degree": degree_name,
                    "university": uni_name,
                    "year": "Unknown" # Year extraction is harder regex
                })
                
    return education

def determine_experience_level(text):
    """Determine if Fresher or Experienced."""
    text_lower = text.lower()
    
    # Signal 1: Headers
    has_experience_header = any(h in text_lower for h in ["work experience", "professional experience", "employment history"])
    
    # Signal 2: Keywords
    is_student = "student" in text_lower or "intern" in text_lower
    
    if not has_experience_header:
        return "Fresher"
        
    # If explicit "Senior" or "Lead" roles found, likely experienced
    if "senior" in text_lower or "lead" in text_lower or "manager" in text_lower:
        return "Experienced"
        
    # Default to Fresher if only internships found? Hard to say without robust parsing.
    # For now, presence of "Work Experience" header -> Experienced
    return "Experienced"

def ingest_cv(file_path):
    path = Path(file_path)
    if not path.exists():
        log(f"Error: File not found: {file_path}")
        return

    log(f"Parsing {file_path}...")
    
    if path.suffix.lower() == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif path.suffix.lower() == '.docx':
        text = extract_text_from_docx(file_path)
    else:
        log("Unsupported file format. Please use PDF or DOCX.")
        return

    if not text:
        log("Failed to extract text.")
        return

    # Extract info
    contact = extract_contact_info(text)
    skills = extract_skills(text)
    education = extract_education(text)
    exp_level = determine_experience_level(text)
    
    log("\n--- Extracted Data ---")
    log(f"Email: {contact['email']}")
    log(f"Phone: {contact['phone']}")
    log(f"Experience Level: {exp_level}")
    log(f"Skills Found: {len(skills)}")
    log(f"Education Entries: {len(education)}")
    
    # Load existing profile
    profile_path = Path(__file__).parent.parent / ".tmp" / "user_profile.json"
    if profile_path.exists():
        with open(profile_path, 'r') as f:
            try:
                profile = json.load(f)
            except:
                profile = {}
    else:
        profile = {}

    # Update profile
    if "personal_info" not in profile: profile["personal_info"] = {}
    if contact['email']: profile["personal_info"]["email"] = contact['email']
    if contact['phone']: profile["personal_info"]["phone"] = contact['phone']
    
    # Merge skills
    # Convert dict skills to list if necessary or handle hybrid
    # Current profile seems to have dict skills, let's normalize to list for this phase
    current_skills_list = []
    if isinstance(profile.get("skills"), dict):
        current_skills_list = profile["skills"].get("technical", [])
        # Extract names if they are objects
        current_skills_list = [s["name"] if isinstance(s, dict) else s for s in current_skills_list]
    elif isinstance(profile.get("skills"), list):
        current_skills_list = profile["skills"]

    current_skills_set = set(current_skills_list)
    current_skills_set.update(skills)
    profile["skills"] = list(current_skills_set)
    
    # Update education (append new ones)
    if "education" not in profile: profile["education"] = []
    
    # Handle if education is dict (unlikely but user_profile.json had objects)
    # The parser returns list of dicts.
    profile["education"].extend(education)
    
    profile["experience_level"] = exp_level
    
    # Save
    with open(profile_path, 'w') as f:
        json.dump(profile, f, indent=2)
        
    log(f"\n✅ Profile updated at {profile_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        log("Usage: python execution/ingest_cv.py <path_to_cv>")
        sys.exit(1)
    
    ingest_cv(sys.argv[1])
